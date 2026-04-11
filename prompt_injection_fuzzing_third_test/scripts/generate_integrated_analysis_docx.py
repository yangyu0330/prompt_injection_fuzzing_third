from __future__ import annotations

import json
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
OUTPUT_DOCX = ROOT / "output" / "doc" / "integrated_analysis_200runs_ko.docx"


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def set_default_font(document: Document, font_name: str = "Malgun Gothic") -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = font_name
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal_style.font.size = Pt(10.5)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(level=level)
    run = heading.add_run(text)
    run.font.size = Pt(15 if level == 1 else 12)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def fmt_rate(v: float) -> str:
    return f"{v:.4f}"


def fmt_ms(v: float) -> str:
    return f"{v:.2f}"


def fmt_bool(v: bool) -> str:
    return "예" if v else "아니오"


def is_pred_positive(row: dict[str, Any]) -> bool:
    return bool(row.get("detected_pre") or row.get("detected_post") or row.get("blocked_effectively"))


def is_attack_success(row: dict[str, Any]) -> bool:
    return bool(row.get("secret_exposed") or row.get("unauthorized_tool_used") or row.get("task_hijacked"))


def layer_counts(scorecard: dict[str, Any]) -> dict[str, int]:
    counts = Counter(str(row.get("layer", "")) for row in scorecard.get("results", []))
    return {k: counts.get(k, 0) for k in ("L1", "L2", "L3")}


def case_level_summary(scorecard: dict[str, Any]) -> dict[tuple[str, str], dict[str, float]]:
    buckets: dict[tuple[str, str], dict[str, float]] = defaultdict(
        lambda: {
            "n": 0.0,
            "pred_positive": 0.0,
            "attack_success": 0.0,
            "blocked": 0.0,
        }
    )
    for row in scorecard.get("results", []):
        key = (str(row.get("layer", "")), str(row.get("case_id", "")))
        bucket = buckets[key]
        bucket["n"] += 1
        bucket["pred_positive"] += 1 if is_pred_positive(row) else 0
        bucket["attack_success"] += 1 if is_attack_success(row) else 0
        bucket["blocked"] += 1 if bool(row.get("blocked_effectively")) else 0
    return buckets


def build_doc() -> Path:
    combined = load_json(REPORTS / "scorecard.json")
    l1 = load_json(REPORTS / "scorecard_l1.json")
    l2 = load_json(REPORTS / "scorecard_l2.json")
    l3 = load_json(REPORTS / "scorecard_l3.json")
    compare = load_json(REPORTS / "ko_en_compare.json")
    cases = case_level_summary(combined)
    layer_run_counts = layer_counts(combined)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_default_font(document)

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("프롬프트 인젝션 퍼징 종합 분석 보고서")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f"L1 / L2 / L3 통합 결과 (총 {combined['run']['total_runs']} runs, "
        f"L1={layer_run_counts['L1']}, L2={layer_run_counts['L2']}, L3={layer_run_counts['L3']})"
    )

    stamp = document.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp.add_run(f"생성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    add_heading(document, "1. 실행 개요", level=1)
    add_table(
        document,
        ["항목", "값"],
        [
            ["총 실행 수", str(combined["run"]["total_runs"])],
            [
                "레이어별 실행 수",
                f"L1={layer_run_counts['L1']}, L2={layer_run_counts['L2']}, L3={layer_run_counts['L3']}",
            ],
            ["템플릿 수", str(combined["run"]["template_count"])],
            ["케이스 수", str(combined["run"]["case_count"])],
            ["Enforcement 모드", "annotate"],
            ["Coverage 통과", fmt_bool(bool(combined["coverage"]["passed"]))],
        ],
    )
    add_paragraph(
        document,
        "참고: 이번 실험은 annotate 모드로 수행되어 탐지 여부를 기록하지만 실제 차단(block)은 수행하지 않습니다.",
    )

    add_heading(document, "2. 레이어별 핵심 지표", level=1)
    add_table(
        document,
        ["레이어", "Recall", "Precision", "FPR", "ASR", "STCR", "평균 지연(ms)", "처리량(rps)"],
        [
            [
                "L1",
                fmt_rate(l1["metrics"]["recall"]),
                fmt_rate(l1["metrics"]["precision"]),
                fmt_rate(l1["metrics"]["fpr"]),
                fmt_rate(l1["metrics"]["asr"]),
                fmt_rate(l1["metrics"]["stcr"]),
                fmt_ms(l1["latency"]["avg_latency_ms"]),
                fmt_rate(l1["latency"]["throughput_rps"]),
            ],
            [
                "L2",
                fmt_rate(l2["metrics"]["recall"]),
                fmt_rate(l2["metrics"]["precision"]),
                fmt_rate(l2["metrics"]["fpr"]),
                fmt_rate(l2["metrics"]["asr"]),
                fmt_rate(l2["metrics"]["stcr"]),
                fmt_ms(l2["latency"]["avg_latency_ms"]),
                fmt_rate(l2["latency"]["throughput_rps"]),
            ],
            [
                "L3",
                fmt_rate(l3["metrics"]["recall"]),
                fmt_rate(l3["metrics"]["precision"]),
                fmt_rate(l3["metrics"]["fpr"]),
                fmt_rate(l3["metrics"]["asr"]),
                fmt_rate(l3["metrics"]["stcr"]),
                fmt_ms(l3["latency"]["avg_latency_ms"]),
                fmt_rate(l3["latency"]["throughput_rps"]),
            ],
            [
                "통합",
                fmt_rate(combined["metrics"]["recall"]),
                fmt_rate(combined["metrics"]["precision"]),
                fmt_rate(combined["metrics"]["fpr"]),
                fmt_rate(combined["metrics"]["asr"]),
                fmt_rate(combined["metrics"]["stcr"]),
                fmt_ms(combined["latency"]["avg_latency_ms"]),
                fmt_rate(combined["latency"]["throughput_rps"]),
            ],
        ],
    )

    add_heading(document, "3. 종합 해석", level=1)
    add_bullet(
        document,
        "L1과 L2의 지표가 사실상 동일합니다. 게이트웨이 래핑(L2)으로 인한 추가 품질 저하는 관찰되지 않았습니다.",
    )
    add_bullet(
        document,
        "L1/L2에서 KO hard 케이스는 지속적으로 미탐지이며, benign 케이스는 과탐지 패턴이 유지됩니다.",
    )
    add_bullet(
        document,
        "L3는 local scenario baseline이라 공격 성공(ASR=1.0)이 높게 나오며, 이는 시나리오 하네스 검증 성격이 강합니다.",
    )
    add_bullet(
        document,
        "통합 지표 해석 시 L1/L2 detector 품질과 L3 outcome 시나리오 결과를 분리해 보는 것이 안전합니다.",
    )

    add_heading(document, "4. KO/EN 비교 분석", level=1)
    add_table(
        document,
        ["비교 축", "값"],
        [
            ["EN clean vs KO clean 점수 차", fmt_rate(compare["baseline_parity"]["score_gap"])],
            ["Baseline 라벨 불일치", fmt_bool(bool(compare["baseline_parity"]["label_mismatch"]))],
            ["KO hard 점수 하락(ko_clean 대비)", fmt_rate(compare["ko_robustness"]["score_drop"])],
            ["KO hard 실패 여부", fmt_bool(bool(compare["ko_robustness"]["hard_mode_failure"]))],
            ["Benign false positive 존재", fmt_bool(bool(compare["benign_false_positive_check"]["benign_false_positive"]))],
        ],
    )
    add_paragraph(
        document,
        "요약: clean 조건의 EN/KO parity는 양호하지만, KO hard robustness와 benign specificity는 개선이 필요합니다.",
    )

    add_heading(document, "5. 케이스별 요약 스냅샷", level=1)
    rows: list[list[str]] = []
    for (layer, case_id), bucket in sorted(cases.items()):
        n = max(int(bucket["n"]), 1)
        rows.append(
            [
                layer,
                case_id,
                str(n),
                fmt_rate(bucket["pred_positive"] / n),
                fmt_rate(bucket["attack_success"] / n),
                fmt_rate(bucket["blocked"] / n),
            ]
        )
    add_table(
        document,
        ["레이어", "케이스 ID", "N", "탐지 양성 비율", "공격 성공 비율", "차단 비율"],
        rows,
    )

    add_heading(document, "6. 리스크 우선순위", level=1)
    add_number(document, "High: L1/L2에서 benign 과탐지 비율이 매우 높아 실사용 영향이 큽니다.")
    add_number(document, "High: KO hard(난독화) 케이스의 탐지 공백이 남아 있습니다.")
    add_number(document, "Medium: annotate-only 결과라 실차단 능력은 검증되지 않았습니다.")
    add_number(document, "Medium: L3는 로컬 시나리오 베이스라인이라 외부 운영 스택 일반화에는 한계가 있습니다.")

    add_heading(document, "7. 권장 후속 조치", level=1)
    add_number(document, "동일 패키지를 block 모드로 재실행해 실차단율과 benign 영향(BOR)을 확인합니다.")
    add_number(document, "KO hard 변형을 jamo 외 spacing/homoglyph/mixed-script로 확장합니다.")
    add_number(document, "quoted/교육/분석 문맥 중심 benign hard-negative를 확장해 specificity를 개선합니다.")
    add_number(document, "L3를 실제 gateway/agent/tool trace 대상으로 연결해 외부 타당성을 확보합니다.")

    add_heading(document, "8. 근거 파일", level=1)
    add_table(
        document,
        ["아티팩트", "경로"],
        [
            ["통합 scorecard", str(REPORTS / "scorecard.json")],
            ["L1 scorecard", str(REPORTS / "scorecard_l1.json")],
            ["L2 scorecard", str(REPORTS / "scorecard_l2.json")],
            ["L3 scorecard", str(REPORTS / "scorecard_l3.json")],
            ["KO/EN 비교", str(REPORTS / "ko_en_compare.json")],
            ["통합 분석 Markdown", str(REPORTS / "integrated_analysis_200runs.md")],
        ],
    )

    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def main() -> None:
    output = build_doc()
    print(output)


if __name__ == "__main__":
    main()

