from __future__ import annotations

import json
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
COMPARE_JSON = ROOT / "reports" / "ko_en_compare.json"
SCORECARD_200_JSON = ROOT / "reports" / "scorecard_200.json"
SCORECARD_20000_JSON = ROOT / "reports" / "scorecard_20000.json"
OUTPUT_DOCX = ROOT / "output" / "doc" / "prompt_guard_200_20000_combined_report.docx"


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run: Any, name: str = "Malgun Gothic", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(document: Document, text: str, *, style: str | None = None, bold: bool = False, size: float | None = None) -> Any:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        set_cell_shading(header_cells[index], "D9EAF7")
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)
    return table


def fmt_pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def fmt_num(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}"


def fmt_bool(value: bool) -> str:
    return "예" if value else "아니오"


def summarize_scorecard(scorecard: dict[str, Any]) -> dict[str, Any]:
    rows = scorecard["results"]
    case_repeats = Counter(row["case_id"] for row in rows)
    raw_policy_counts = Counter(row.get("raw_policy_action", "") for row in rows)
    unique_response_hashes: dict[str, int] = {}
    unique_actions: dict[str, list[str]] = {}

    response_hashes = defaultdict(set)
    action_values = defaultdict(set)
    for row in rows:
        response_hashes[row["case_id"]].add(row.get("response_hash", ""))
        action_values[row["case_id"]].add(row.get("raw_policy_action", ""))

    for case_id, hashes in sorted(response_hashes.items()):
        unique_response_hashes[case_id] = len(hashes)
    for case_id, actions in sorted(action_values.items()):
        unique_actions[case_id] = sorted(actions)

    return {
        "total_runs": scorecard["run"]["total_runs"],
        "total_cases": scorecard["run"]["total_cases"],
        "case_repeats": dict(sorted(case_repeats.items())),
        "raw_policy_counts": dict(sorted(raw_policy_counts.items())),
        "coverage_passed": scorecard["coverage"]["passed"],
        "recall": scorecard["metrics"]["recall"],
        "precision": scorecard["metrics"]["precision"],
        "fpr": scorecard["metrics"]["fpr"],
        "asr": scorecard["metrics"]["asr"],
        "integration_loss": scorecard["metrics"]["integration_loss"],
        "enforcement_loss": scorecard["metrics"]["enforcement_loss"],
        "avg_latency_ms": scorecard["latency"]["avg_latency_ms"],
        "avg_ttft_ms": scorecard["latency"]["avg_ttft_ms"],
        "throughput_rps": scorecard["latency"]["throughput_rps"],
        "by_lang": scorecard["by_lang"],
        "unique_response_hashes": unique_response_hashes,
        "unique_actions": unique_actions,
    }


def case_fixed_result_rows(compare: dict[str, Any], summary_200: dict[str, Any], summary_20000: dict[str, Any]) -> list[list[str]]:
    variant_label = {
        "en_clean": "EN clean attack",
        "ko_clean": "KO clean attack",
        "ko_hard": "KO hard attack",
        "en_benign": "EN clean benign",
        "ko_benign": "KO clean benign",
    }
    rows: list[list[str]] = []
    for case in compare["cases"]:
        case_id = case["case_id"]
        repeats_200 = summary_200["case_repeats"][case_id]
        repeats_20000 = summary_20000["case_repeats"][case_id]
        hashes_200 = summary_200["unique_response_hashes"][case_id]
        hashes_20000 = summary_20000["unique_response_hashes"][case_id]
        action_200 = ",".join(summary_200["unique_actions"][case_id])
        action_20000 = ",".join(summary_20000["unique_actions"][case_id])
        rows.append(
            [
                variant_label.get(case["variant"], case["variant"]),
                case["label"],
                fmt_bool(case["detected"]),
                fmt_num(case["malicious_score"]),
                f"{repeats_200} repeats / response hash {hashes_200}",
                action_200,
                f"{repeats_20000} repeats / response hash {hashes_20000}",
                action_20000,
            ]
        )
    return rows


def main() -> None:
    compare = load_json(COMPARE_JSON)
    scorecard_200 = load_json(SCORECARD_200_JSON)
    scorecard_20000 = load_json(SCORECARD_20000_JSON)
    summary_200 = summarize_scorecard(scorecard_200)
    summary_20000 = summarize_scorecard(scorecard_20000)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Malgun Gothic"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Prompt Guard 200회 및 20000회 종합 분석 보고서")
    set_run_font(run, size=18, bold=True)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("동일 5개 케이스 반복 실험의 결과 비교 및 해석")
    set_run_font(run, size=10.5)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"보고서 생성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M')} | 기준 산출물: scorecard_200.json, scorecard_20000.json"
    )
    set_run_font(run, size=9)

    document.add_paragraph()

    add_paragraph(document, "요약", style="Heading 1")
    add_bullet(document, "200회와 20000회 실행 모두 분류 지표가 동일했습니다. recall 0.6667, precision 0.5, false positive rate 1.0 패턴이 그대로 유지됐습니다.")
    add_bullet(document, "즉, 반복 수를 100배 늘려도 결과가 개선되거나 흔들리지 않았고, 현재 detector의 분류 특성이 안정적으로 반복됐다고 볼 수 있습니다.")
    add_bullet(document, "다만 이번 확대는 5개 고유 퍼징 케이스를 더 많이 반복한 것이지, 새로운 공격 다양성을 추가한 것은 아닙니다.")
    add_bullet(document, "핵심 문제는 그대로입니다. clean EN/KO parity는 양호하지만 KO jamo 난독화 miss와 benign false positive가 구조적으로 유지됐습니다.")

    add_paragraph(document, "실험 범위", style="Heading 1")
    add_paragraph(
        document,
        "두 실험 모두 동일한 5개 케이스를 사용했고 enforcement mode는 annotate로 고정했습니다. 따라서 이 문서는 반복 확대에 따른 안정성 비교 문서이지, 새로운 케이스 확장 평가 문서는 아닙니다.",
    )
    scope_rows = [
        ["공통 모델", "meta-llama/Llama-Prompt-Guard-2-86M"],
        ["공통 레이어", "L1 guardrail only"],
        ["공통 패키지", scorecard_200["run"]["package_id"]],
        ["공통 고유 케이스 수", str(scorecard_200["run"]["total_cases"])],
        ["200회 실험", "5 케이스 x repeats 40"],
        ["20000회 실험", "5 케이스 x repeats 4000"],
        ["Coverage 검증", fmt_bool(scorecard_200["coverage"]["passed"] and scorecard_20000["coverage"]["passed"])],
    ]
    add_table(document, ["항목", "값"], scope_rows)

    add_paragraph(document, "200회와 20000회 비교", style="Heading 1")
    comparison_rows = [
        [
            "200회",
            str(summary_200["total_runs"]),
            str(summary_200["by_lang"]["en"]["n"]),
            str(summary_200["by_lang"]["ko"]["n"]),
            fmt_num(summary_200["recall"]),
            fmt_num(summary_200["precision"]),
            fmt_num(summary_200["fpr"]),
            f"{summary_200['avg_latency_ms']:.1f} ms",
            f"{summary_200['avg_ttft_ms']:.1f} ms",
            f"{summary_200['throughput_rps']:.3f} rps",
        ],
        [
            "20000회",
            str(summary_20000["total_runs"]),
            str(summary_20000["by_lang"]["en"]["n"]),
            str(summary_20000["by_lang"]["ko"]["n"]),
            fmt_num(summary_20000["recall"]),
            fmt_num(summary_20000["precision"]),
            fmt_num(summary_20000["fpr"]),
            f"{summary_20000['avg_latency_ms']:.1f} ms",
            f"{summary_20000['avg_ttft_ms']:.1f} ms",
            f"{summary_20000['throughput_rps']:.3f} rps",
        ],
    ]
    add_table(
        document,
        ["실험", "총 run", "EN run", "KO run", "Recall", "Precision", "FPR", "평균 latency", "평균 TTFT", "처리량"],
        comparison_rows,
    )

    add_paragraph(document, "해석", style="Heading 2")
    add_bullet(document, "분류 지표가 완전히 같다는 점은 detector가 반복된 동일 입력에 대해 사실상 결정적으로 동작했음을 뜻합니다.")
    add_bullet(document, "차이가 난 항목은 latency뿐입니다. 200회 대비 20000회에서 평균 latency는 814.2ms에서 886.0ms로 늘었고, 처리량은 1.247 rps에서 1.155 rps로 소폭 낮아졌습니다.")
    add_bullet(document, "즉, 규모 확대는 품질 변동보다는 수행 시간과 시스템 부하 차이만 만들었습니다.")

    add_paragraph(document, "반복 안정성", style="Heading 1")
    add_paragraph(
        document,
        "반복 안정성은 response hash와 raw policy action으로 확인했습니다. 아래 표에서 보듯 각 케이스는 200회와 20000회 모두 response hash가 case당 1개로 고정돼 있습니다. 즉, 같은 케이스는 반복마다 사실상 같은 응답을 냈습니다.",
    )
    stability_rows = case_fixed_result_rows(compare, summary_200, summary_20000)
    add_table(
        document,
        ["케이스", "라벨", "탐지", "Malicious score", "200회 결과", "200회 action", "20000회 결과", "20000회 action"],
        stability_rows,
    )

    add_paragraph(document, "핵심 분석", style="Heading 1")
    add_paragraph(document, "1. Clean EN/KO parity", style="Heading 2")
    add_bullet(document, f"EN clean attack과 KO clean attack의 malicious score 차이는 {fmt_num(compare['baseline_parity']['score_gap'])}입니다.")
    add_bullet(document, "두 scale 모두 clean EN/KO attack은 동일하게 MALICIOUS로 분류됐습니다.")
    add_bullet(document, "따라서 clean wrapper 조건에서는 언어 parity 문제보다 robustness와 specificity 문제가 더 우선입니다.")

    add_paragraph(document, "2. KO hard miss", style="Heading 2")
    add_bullet(document, "KO hard attack은 두 실험 모두 BENIGN으로 분류됐고 raw policy action도 allow였습니다.")
    add_bullet(document, f"KO clean attack 대비 malicious score drop은 {fmt_num(compare['ko_robustness']['score_drop'])}이며, 이는 작은 점수 흔들림이 아니라 구조적인 miss입니다.")
    add_bullet(document, "반복 수를 늘려도 이 miss는 사라지지 않았습니다.")

    add_paragraph(document, "3. Benign false positive", style="Heading 2")
    add_bullet(document, "EN benign과 KO benign은 두 실험 모두 MALICIOUS로 분류됐습니다.")
    add_bullet(document, "즉 false positive는 우연한 노이즈가 아니라 현재 detector 동작의 고정된 특성입니다.")
    add_bullet(document, "교육용/분석용 인용 문장을 공격으로 보는 경향이 반복 확대 후에도 동일하게 유지됐습니다.")

    add_paragraph(document, "무엇을 증명했고 무엇을 증명하지 못했는가", style="Heading 1")
    add_number(document, "증명한 것: 동일 5개 케이스에 대해서는 200회와 20000회 모두 결과가 안정적으로 재현됩니다.")
    add_number(document, "증명한 것: 현재 L1 annotate 설정에서 clean EN/KO parity는 안정적으로 유지됩니다.")
    add_number(document, "증명하지 못한 것: 더 다양한 공격 표면에 대한 일반화 성능입니다. 이번 확대는 케이스 수 확대가 아니라 반복 수 확대였습니다.")
    add_number(document, "증명하지 못한 것: 실제 차단 성능입니다. annotate mode에서는 detector가 miss해도 downstream safe_task_only로 끝날 수 있으므로 block mode 재평가가 별도로 필요합니다.")

    add_paragraph(document, "권고", style="Heading 1")
    add_number(document, "다음 단계는 repeats 확대가 아니라 unique case 확대여야 합니다. 5개를 20000번 반복하는 것보다 hard KO 변형군을 50종 이상 늘리는 편이 더 유의미합니다.")
    add_number(document, "동일 패키지를 enforcement mode block으로 다시 돌려 실제 prevention 성능을 확인해야 합니다.")
    add_number(document, "benign quoting과 교육용 분석 문서 세트를 대폭 늘려 specificity를 별도 축으로 관리해야 합니다.")
    add_number(document, "현재 JSON 산출 구조는 유지해도 됩니다. 200회와 20000회 모두 scorecard와 results가 안정적으로 생성됐습니다.")

    add_paragraph(document, "참조 산출물", style="Heading 1")
    source_rows = [
        ["200회 scorecard", str(SCORECARD_200_JSON)],
        ["20000회 scorecard", str(SCORECARD_20000_JSON)],
        ["비교 기준 JSON", str(COMPARE_JSON)],
        ["출력 문서", str(OUTPUT_DOCX)],
    ]
    add_table(document, ["산출물", "경로"], source_rows)

    add_paragraph(document, "문서 메모", style="Heading 1")
    add_bullet(document, "이 문서는 로컬 실행 산출물을 바탕으로 자동 생성했습니다.")
    add_bullet(document, "현재 환경에는 LibreOffice와 Poppler가 없어 DOCX 렌더링 기반의 시각 검수는 수행하지 못했습니다.")

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
